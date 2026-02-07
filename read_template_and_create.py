#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
读取模板文件并创建CreateNow项目的软件著作权登记申请表
完全按照模板格式，使用上海首界科技的申请人信息
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def read_template_structure():
    """读取并分析模板文件结构"""
    template_path = r'著作权范例\01-【新】软件著作权登记申请表-模板(2).docx'

    print("=" * 80)
    print("读取模板文件结构")
    print("=" * 80)

    try:
        doc = Document(template_path)

        # 打印所有段落
        print("\n【段落内容】")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"段落 {i}: {para.text}")

        # 打印所有表格
        print("\n【表格内容】")
        for table_idx, table in enumerate(doc.tables):
            print(f"\n表格 {table_idx + 1}:")
            print(f"  行数: {len(table.rows)}, 列数: {len(table.columns)}")
            for row_idx, row in enumerate(table.rows):
                cells_text = []
                for cell in row.cells:
                    cells_text.append(cell.text.strip())
                print(f"  行 {row_idx}: {cells_text}")

        return doc

    except Exception as e:
        print(f"错误: 无法读取模板文件: {e}")
        import traceback
        traceback.print_exc()
        return None

def create_copyright_form_from_template():
    """基于模板创建CreateNow的申请表"""

    # 先读取模板了解结构
    template_doc = read_template_structure()

    if template_doc is None:
        print("\n无法读取模板，退出")
        return None

    print("\n" + "=" * 80)
    print("开始创建CreateNow申请表")
    print("=" * 80)

    # 加载模板作为基础
    template_path = r'著作权范例\01-【新】软件著作权登记申请表-模板(2).docx'
    doc = Document(template_path)

    # 根据模板结构填充CreateNow的信息
    # 这里需要根据实际模板结构来修改对应的表格单元格

    # 假设第一个表格是软件基本信息
    if len(doc.tables) > 0:
        table = doc.tables[0]
        print(f"\n修改第1个表格（软件基本信息）")

        # 根据模板的实际结构来填充
        # 这里需要查看模板后确定具体的行列位置

    # 保存新文档
    output_path = 'CreateNow软件著作权登记申请表-正式版.docx'
    doc.save(output_path)
    print(f"\n✓ 文档已生成: {output_path}")

    return output_path

if __name__ == "__main__":
    try:
        create_copyright_form_from_template()
    except Exception as e:
        print(f"\n错误: {e}")
        import traceback
        traceback.print_exc()
