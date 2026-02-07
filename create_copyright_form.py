#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成CreateNow项目的软件著作权登记申请表
基于中国软件著作权登记标准格式
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    print("成功导入 python-docx 库")
except ImportError as e:
    print(f"错误: 无法导入 python-docx 库: {e}")
    print("请运行: pip install python-docx")
    exit(1)

def create_copyright_application():
    """创建软件著作权登记申请表"""

    print("开始创建文档...")
    doc = Document()

    # 设置文档默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('计算机软件著作权登记申请表')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = '黑体'

    doc.add_paragraph()  # 空行

    # 软件基本信息
    doc.add_paragraph('一、软件基本信息').runs[0].font.bold = True

    table1 = doc.add_table(rows=10, cols=2)
    table1.style = 'Table Grid'

    # 填充表格内容
    rows_data = [
        ('软件全称', 'CreateNow AI短视频生成平台'),
        ('软件简称', 'CreateNow'),
        ('软件版本号', 'V1.0'),
        ('开发完成日期', '2025年1月'),
        ('首次发表日期', '2025年1月'),
        ('软件开发情况', '独立开发'),
        ('原始取得权利方式', '原始开发'),
        ('权利范围', '全部权利'),
        ('软件用途和技术特点', 'AI驱动的短视频内容创作平台，通过对话式交互实现剧本创作、资产管理、分镜设计和视频生成的全流程自动化。采用FastAPI后端、React前端架构，支持多种AI模型接入（OpenAI、阿里云等），提供文生图、图生视频等AIGC能力。'),
        ('编程语言', 'Python, TypeScript, JavaScript')
    ]

    for i, (label, value) in enumerate(rows_data):
        table1.rows[i].cells[0].text = label
        table1.rows[i].cells[1].text = value
        # 设置第一列宽度
        table1.rows[i].cells[0].width = Inches(2.0)

    doc.add_paragraph()  # 空行

    # 著作权人信息
    doc.add_paragraph('二、著作权人信息').runs[0].font.bold = True

    table2 = doc.add_table(rows=8, cols=2)
    table2.style = 'Table Grid'

    author_data = [
        ('著作权人性质', '个人'),
        ('姓名/名称', '[请填写著作权人姓名]'),
        ('国籍/地区', '中国'),
        ('证件类型', '身份证'),
        ('证件号码', '[请填写身份证号]'),
        ('联系电话', '[请填写联系电话]'),
        ('电子邮箱', '[请填写电子邮箱]'),
        ('详细地址', '[请填写详细地址]')
    ]

    for i, (label, value) in enumerate(author_data):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = value
        table2.rows[i].cells[0].width = Inches(2.0)

    doc.add_paragraph()  # 空行

    # 软件鉴别材料
    doc.add_paragraph('三、软件鉴别材料').runs[0].font.bold = True

    table3 = doc.add_table(rows=3, cols=2)
    table3.style = 'Table Grid'

    material_data = [
        ('程序鉴别材料', '源程序代码前30页和后30页，共60页（每页不少于50行）'),
        ('文档鉴别材料', '软件操作说明书（图文并茂，不少于10页）'),
        ('例外交存', '无')
    ]

    for i, (label, value) in enumerate(material_data):
        table3.rows[i].cells[0].text = label
        table3.rows[i].cells[1].text = value
        table3.rows[i].cells[0].width = Inches(2.0)

    doc.add_paragraph()  # 空行

    # 软件功能和技术特点详细说明
    doc.add_paragraph('四、软件功能和技术特点详细说明').runs[0].font.bold = True

    features = doc.add_paragraph()
    features.add_run('1. 核心功能：\n').font.bold = True
    features.add_run(
        '• 对话式剧本创作：通过自然语言对话，AI自动提取角色、场景、道具等资产\n'
        '• 智能资产管理：支持资产继承、变体创建、批量生成图片\n'
        '• 分镜设计：可视化分镜编辑，支持拖拽排序、批量操作\n'
        '• 视频生成：基于分镜图片生成视频，支持多种AI模型\n'
        '• 项目管理：多项目支持，独立配置AI接口\n\n'
    )

    features.add_run('2. 技术特点：\n').font.bold = True
    features.add_run(
        '• 前后端分离架构：FastAPI + React，支持开发/生产双模式部署\n'
        '• 文件存储系统：无数据库设计，JSON文件存储，便于迁移和备份\n'
        '• 流式响应：WebSocket实时传输AI生成内容\n'
        '• 多模型兼容：支持OpenAI、阿里云DashScope等多种AI服务\n'
        '• 模块化设计：清晰的服务层、API层、模型层分离\n'
        '• 跨平台支持：提供Windows批处理和Shell脚本启动方式\n\n'
    )

    features.add_run('3. 技术栈：\n').font.bold = True
    features.add_run(
        '• 后端：Python 3.10+, FastAPI, Pydantic, Uvicorn, httpx\n'
        '• 前端：React 18, TypeScript, Vite, Zustand, Tailwind CSS, Axios\n'
        '• AI集成：OpenAI API, 阿里云DashScope API\n'
        '• 开发工具：ESLint, TypeScript Compiler, Hot Reload\n'
    )

    doc.add_paragraph()  # 空行

    # 申请人声明
    doc.add_paragraph('五、申请人声明').runs[0].font.bold = True

    declaration = doc.add_paragraph()
    declaration.add_run(
        '本人郑重声明：\n'
        '1. 本申请表中所填写的内容真实、准确、完整；\n'
        '2. 本软件系原始开发，不存在侵犯他人著作权的情形；\n'
        '3. 本软件未进行过著作权登记；\n'
        '4. 本人承诺对以上声明承担法律责任。\n'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    signature = doc.add_paragraph()
    signature.add_run('申请人签名：________________    日期：________________')

    # 保存文档
    output_path = 'CreateNow软件著作权登记申请表.docx'
    doc.save(output_path)
    print(f"✓ 文档已成功生成: {output_path}")
    print("\n请注意：")
    print("1. 需要填写著作权人的个人信息（姓名、身份证号、联系方式、地址）")
    print("2. 需要确认开发完成日期和首次发表日期")
    print("3. 需要准备源代码文件（前30页+后30页）")
    print("4. 需要准备软件操作说明书（不少于10页）")

    return output_path

if __name__ == "__main__":
    try:
        create_copyright_application()
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
