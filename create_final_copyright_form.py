#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
基于模板创建CreateNow项目的软件著作权登记申请表
使用上海首界科技的申请人信息
"""

from docx import Document
import os

def create_createnow_copyright_form():
    """创建CreateNow的软件著作权登记申请表"""

    template_path = r'著作权范例\01-【新】软件著作权登记申请表-模板(2).docx'

    print("正在加载模板...")
    doc = Document(template_path)

    # 模板有1个主表格，包含38行5列
    # 我们需要修改表格中的具体内容

    if len(doc.tables) > 0:
        table = doc.tables[0]
        print(f"找到表格，共 {len(table.rows)} 行")

        # 根据模板结构，修改对应的单元格内容
        # 需要找到每个字段所在的行，然后修改对应的值

        # 遍历所有行，找到需要修改的字段
        for row_idx, row in enumerate(table.rows):
            cells = row.cells

            # 获取第一列的文本作为字段名
            if len(cells) > 0:
                field_name = cells[0].text.strip()

                # 根据字段名修改对应的值
                if '软件名称' in field_name and '全称' in field_name:
                    # 软件全称
                    if len(cells) > 1:
                        cells[1].text = 'CreateNow AI短视频生成平台'
                        print(f"行 {row_idx}: 设置软件全称")

                elif '软件名称' in field_name and '简称' in field_name:
                    # 软件简称
                    if len(cells) > 1:
                        cells[1].text = 'CreateNow'
                        print(f"行 {row_idx}: 设置软件简称")

                elif '分类号' in field_name:
                    # 分类号
                    if len(cells) > 1:
                        cells[1].text = 'TP391'
                        print(f"行 {row_idx}: 设置分类号")

                elif '版本号' in field_name:
                    # 版本号
                    if len(cells) > 1:
                        cells[1].text = 'V1.0'
                        print(f"行 {row_idx}: 设置版本号")

                elif '软件作品说明' in field_name:
                    # 软件作品说明
                    if len(cells) > 1:
                        cells[1].text = ('CreateNow是一款基于人工智能技术的短视频内容创作平台。'
                                       '该软件通过对话式交互方式，帮助用户完成从剧本创作、资产管理、'
                                       '分镜设计到视频生成的全流程自动化创作。系统采用前后端分离架构，'
                                       '后端使用Python FastAPI框架，前端使用React+TypeScript技术栈，'
                                       '集成多种AI模型（OpenAI、阿里云等），提供文生图、图生视频等AIGC能力。'
                                       '软件支持多项目管理、资产继承与变体、实时流式响应等特性，'
                                       '采用文件存储系统，无需数据库，便于部署和迁移。')
                        print(f"行 {row_idx}: 设置软件作品说明")

                elif '开发完成日期' in field_name:
                    # 开发完成日期
                    if len(cells) > 1:
                        cells[1].text = '2025年1月20日'
                        print(f"行 {row_idx}: 设置开发完成日期")

                elif '发表日期' in field_name or '首次发表日期' in field_name:
                    # 发表日期
                    if len(cells) > 1:
                        cells[1].text = '2025年1月28日'
                        print(f"行 {row_idx}: 设置发表日期")

                elif '开发方式' in field_name:
                    # 开发方式
                    if len(cells) > 1:
                        cells[1].text = '独立开发'
                        print(f"行 {row_idx}: 设置开发方式")

                elif '硬件环境' in field_name:
                    # 硬件环境
                    if len(cells) > 1:
                        cells[1].text = 'Intel Core i5及以上处理器，8GB及以上内存，100GB及以上硬盘空间'
                        print(f"行 {row_idx}: 设置硬件环境")

                elif '软件环境' in field_name or '操作系统' in field_name:
                    # 软件环境/操作系统
                    if len(cells) > 1:
                        cells[1].text = 'Windows 10/11, macOS 10.15+, Linux (Ubuntu 20.04+)'
                        print(f"行 {row_idx}: 设置软件环境")

                elif '编程语言' in field_name:
                    # 编程语言
                    if len(cells) > 1:
                        cells[1].text = 'Python 3.10+, TypeScript, JavaScript'
                        print(f"行 {row_idx}: 设置编程语言")

                elif '源程序量' in field_name or '代码行数' in field_name:
                    # 源程序量
                    if len(cells) > 1:
                        cells[1].text = '约15000行'
                        print(f"行 {row_idx}: 设置源程序量")

                elif '主要功能' in field_name or '功能和技术特点' in field_name:
                    # 主要功能和技术特点
                    if len(cells) > 1:
                        cells[1].text = ('1.对话式剧本创作：通过自然语言对话自动提取角色、场景、道具等资产；'
                                       '2.智能资产管理：支持资产继承、变体创建、批量图片生成；'
                                       '3.分镜设计：可视化分镜编辑，支持拖拽排序、批量操作；'
                                       '4.视频生成：基于分镜图片生成视频，支持多种AI模型；'
                                       '5.项目管理：多项目支持，独立配置AI接口；'
                                       '6.技术特点：前后端分离、WebSocket流式响应、文件存储、多模型兼容。')
                        print(f"行 {row_idx}: 设置主要功能")

                elif '著作权人' in field_name and '名称' in field_name:
                    # 著作权人名称 - 使用上海首界科技
                    if len(cells) > 1:
                        cells[1].text = '上海首界科技有限公司'
                        print(f"行 {row_idx}: 设置著作权人名称")

                elif '国籍' in field_name or '国家/地区' in field_name:
                    # 国籍
                    if len(cells) > 1:
                        cells[1].text = '中国'
                        print(f"行 {row_idx}: 设置国籍")

                elif '统一社会信用代码' in field_name or '证件号码' in field_name:
                    # 统一社会信用代码
                    if len(cells) > 1:
                        cells[1].text = '91310115MAD3LX2P53'
                        print(f"行 {row_idx}: 设置统一社会信用代码")

                elif '著作权人类别' in field_name:
                    # 著作权人类别
                    if len(cells) > 1:
                        cells[1].text = '企业法人'
                        print(f"行 {row_idx}: 设置著作权人类别")

    # 保存文档
    output_path = 'CreateNow软件著作权登记申请表-正式版.docx'
    doc.save(output_path)
    print(f"\n✓ 文档已成功生成: {output_path}")
    print("\n请注意：")
    print("1. 已使用上海首界科技有限公司作为著作权人")
    print("2. 已填写CreateNow项目的基本信息")
    print("3. 请检查并补充其他可能需要的详细信息")
    print("4. 需要准备源代码文件（前30页+后30页）和操作说明书")

    return output_path

if __name__ == "__main__":
    try:
        create_createnow_copyright_form()
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
