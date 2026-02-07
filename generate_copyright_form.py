#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成CreateNow项目的软件著作权登记申请表
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def read_template():
    """读取模板文件内容"""
    template_path = r'著作权范例\01-【新】软件著作权登记申请表-模板(2).docx'
    try:
        doc = Document(template_path)
        print("=== 模板内容 ===")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"{i}: {para.text}")

        # 打印表格内容
        print("\n=== 表格内容 ===")
        for table_idx, table in enumerate(doc.tables):
            print(f"\n表格 {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                row_text = [cell.text.strip() for cell in row.cells]
                print(f"  行 {row_idx}: {row_text}")

        return doc
    except Exception as e:
        print(f"读取模板失败: {e}")
        return None

def create_copyright_form():
    """创建CreateNow项目的软件著作权登记申请表"""

    # 先读取模板了解结构
    template_doc = read_template()

    if template_doc is None:
        print("无法读取模板，将创建基础表格")
        doc = Document()
    else:
        # 使用模板作为基础
        doc = template_doc

    # 保存文件
    output_path = 'CreateNow软件著作权登记申请表.docx'
    doc.save(output_path)
    print(f"\n文件已生成: {output_path}")

    return output_path

if __name__ == "__main__":
    create_copyright_form()
